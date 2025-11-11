#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Create a nicely formatted Word (.docx) publication summary from the
Excel export created by `resulam_book_processor.py`.

Output: Book_Exports/books_publication_summary.docx

Requirements: python-docx (install into your venv: pip install python-docx)
"""
import datetime
from pathlib import Path
from typing import Dict

import pandas as pd
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph. Returns the created hyperlink element.

    This uses the low-level OXML API because python-docx doesn't expose a
    direct high-level hyperlink builder.
    """
    # Create a relationship id for this external hyperlink
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and set the relationship id
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')
    # Create a rPr element
    rPr = OxmlElement('w:rPr')
    # Apply the built-in Hyperlink style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)

    # Create the w:t element and set text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)

    # Append the hyperlink to the paragraph XML
    paragraph._p.append(hyperlink)
    return hyperlink


def make_docx_from_excel(excel_path: Path, out_docx: Path) -> None:
    print(f"Loading workbook: {excel_path}")
    sheets: Dict[str, pd.DataFrame] = pd.read_excel(excel_path, sheet_name=None)

    doc = Document()

    # Title page
    title = doc.add_heading('Resulam Books — Publication Summary', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f"Generated: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph('')

    # Add a short intro
    intro = doc.add_paragraph()
    intro.add_run('This document contains a formatted summary of Resulam publications, grouped by language. Click the links to open the printed product page.').italic = True
    doc.add_paragraph('')

    # Style setup: slightly larger font for headings
    for sheet_name, df in sheets.items():
        # Add language heading
        doc.add_heading(sheet_name, level=1)

        if df.empty:
            doc.add_paragraph('No entries for this language.')
            continue

        # Ensure rows are iterated in the order present in the sheet
        for idx, row in df.reset_index(drop=True).iterrows():
            p = doc.add_paragraph()
            # Number and tab
            num_run = p.add_run(f"{idx+1}-\t")
            num_run.bold = True

            # Authors
            authors = row.get('authors') if pd.notna(row.get('authors')) else ''
            if authors:
                p.add_run(f"{authors}, ")

            # Title (bold)
            title_text = row.get('title') if pd.notna(row.get('title')) else 'Untitled'
            title_run = p.add_run(f'"{title_text}"')
            title_run.bold = True
            p.add_run(', ')

            # Pages (prefer pages_paperback -> pages_hard_cover -> pages_ebook -> number_of_pages)
            pages = None
            for col in ['pages_paperback', 'pages_hard_cover', 'pages_ebook', 'number_of_pages']:
                if col in row and pd.notna(row.get(col)) and str(row.get(col)).strip() != '':
                    raw_pages = row.get(col)
                    # Normalize to integer when possible (strip .0)
                    try:
                        pages_int = int(float(raw_pages))
                        pages = str(pages_int)
                    except Exception:
                        # Fallback: extract first integer substring
                        m = re.search(r"(\d+)", str(raw_pages))
                        if m:
                            pages = m.group(1)
                        else:
                            pages = str(raw_pages).strip()
                    break
            if pages:
                p.add_run(f"({pages} pages), ")

            # ISBN candidates
            isbn = None
            for c in ['paperback_isbn13', 'isbn_paperback', 'paperback_isbn', 'paperback_isbn10']:
                if c in row and pd.notna(row.get(c)) and str(row.get(c)).strip() != '':
                    isbn = str(row.get(c)).strip()
                    break
            if isbn:
                p.add_run(f"ISBN-13: {isbn}, ")

            # Publication date
            pub_date = row.get('publication_date')
            if pd.notna(pub_date):
                try:
                    date_str = pd.to_datetime(pub_date).strftime('%B %d, %Y')
                except Exception:
                    date_str = str(pub_date)
                p.add_run(f"{date_str}. ")

            # Printed link as clickable hyperlink
            printed_link = row.get('paperback') if 'paperback' in row else None
            if pd.notna(printed_link) and printed_link:
                p.add_run('Link (Printed): ')
                add_hyperlink(p, printed_link, printed_link)

    # Save document
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_docx)
        print(f"Saved Word document to: {out_docx}")
    except PermissionError:
        # File is probably open in Word. Try a fallback filename.
        fallback = out_docx.with_name(out_docx.stem + '_new' + out_docx.suffix)
        doc.save(fallback)
        print(f"Could not write to {out_docx} (permission denied). Saved to: {fallback}")


if __name__ == '__main__':
    excel = Path('Book_Exports') / '03_all_books_by_language.xlsx'
    out = Path('Book_Exports') / 'books_publication_summary.docx'
    make_docx_from_excel(excel, out)
