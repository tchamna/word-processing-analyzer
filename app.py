import streamlit as st
from docx import Document
from docx.shared import RGBColor
import re
import io
import csv
import zipfile
import xml.etree.ElementTree as ET

def load_document_ignoring_bad_media(docx_bytes: bytes) -> Document:
    """Load a DOCX while gracefully skipping corrupted media.

    Rebuilds the DOCX zip in-memory, excluding corrupted media entries
    and cleaning up media references in relationship files.
    """
    src_bio = io.BytesIO(docx_bytes)
    try:
        with zipfile.ZipFile(src_bio, "r") as src_zip:
            # First pass: identify which media files are corrupted or missing
            corrupt_media = set()
            for name in src_zip.namelist():
                if name.startswith("word/media/"):
                    try:
                        # Test if we can read it
                        src_zip.read(name)
                    except Exception:
                        # Mark as corrupt
                        corrupt_media.add(name)
            
            # Second pass: rebuild ZIP without corrupt media and clean relationships
            new_bio = io.BytesIO()
            with zipfile.ZipFile(new_bio, "w", compression=zipfile.ZIP_DEFLATED) as dst_zip:
                for name in src_zip.namelist():
                    # Skip all media files (corrupt or not) to avoid issues
                    if name.startswith("word/media/"):
                        continue
                    
                    # For .rels files, clean up media references
                    if name.endswith(".rels"):
                        try:
                            data = src_zip.read(name)
                            # Parse and clean XML
                            root = ET.fromstring(data)
                            ns = {'rel': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                            
                            # Remove relationships pointing to media
                            for rel in root.findall('.//rel:Relationship', ns):
                                target = rel.get('Target', '')
                                if 'media/' in target:
                                    root.remove(rel)
                            
                            # Write cleaned XML
                            cleaned_data = ET.tostring(root, encoding='utf-8', xml_declaration=True)
                            dst_zip.writestr(name, cleaned_data)
                        except Exception:
                            # If cleaning fails, write original
                            try:
                                data = src_zip.read(name)
                                dst_zip.writestr(name, data)
                            except Exception:
                                continue
                    else:
                        # Copy other files as-is
                        try:
                            data = src_zip.read(name)
                            dst_zip.writestr(name, data)
                        except Exception:
                            continue
            
            new_bio.seek(0)
            return Document(new_bio)
    except Exception:
        # As a fallback, try loading the original; if it fails, let caller catch
        return Document(io.BytesIO(docx_bytes))

st.set_page_config(
    page_title="Word Format Issue Detector",
    page_icon="📝",
    layout="wide"
)

st.title("Word Format Issue Detector")
st.write("Upload a .docx file to check for formatting issues including colored text, capitalization, punctuation, and spacing.")

# Punctuations that should end a sentence.
PUNCTUATION = ".?!;"
# Punctuations that should not have a space before them.
SPACED_PUNCTUATION = ",."

def check_paragraph(i, para):
    """Check one paragraph for capitalization, punctuation, and color."""
    issues = []
    text = para.text.strip()
    if not text:
        return issues

    # Color check (detect any colored text except black)
    colored_text = ""
    for run in para.runs:
        if run.font.color and run.font.color.rgb and run.font.color.rgb != RGBColor(0, 0, 0):
            colored_text += run.text
    if colored_text:
        issues.append({
            "line": i + 1,
            "type": "Contains colored text",
            "sentence": text,
            "suggestion": f"Review colored text: '{colored_text}'"
        })

    # Capitalization check
    if not text[0].isupper():
        issues.append({
            "line": i + 1,
            "type": "Missing capitalization",
            "sentence": text,
            "suggestion": f"Capitalize the first letter: '{text[0].upper() + text[1:]}'"
        })

    # Punctuation check
    if text[-1] not in PUNCTUATION:
        issues.append({
            "line": i + 1,
            "type": "Missing punctuation",
            "sentence": text,
            "suggestion": f"Add punctuation at the end, e.g., '{text + '.'}'"
        })

    # Spacing before punctuation check
    punct_issues = []
    if re.search(r'\s,', text):
        punct_issues.append("comma")
    if re.search(r'\s\.', text):
        punct_issues.append("period")
    if punct_issues:
        punct_str = " and ".join(punct_issues) if len(punct_issues) > 1 else punct_issues[0]
        corrected = re.sub(r'\s([,.])', r'\1', text)
        issues.append({
            "line": i + 1,
            "type": "Incorrect spacing before punctuation",
            "sentence": text,
            "suggestion": f"Remove spaces before {punct_str}: '{corrected}'"
        })

    return issues

st.title("Docx Formatting Analyzer")
st.write("Upload a .docx file to check for formatting issues including colored text, capitalization, punctuation, and spacing.")

uploaded_file = st.file_uploader("Choose a .docx file", type=["docx"])

if uploaded_file is not None:
    try:
        file_bytes = uploaded_file.read()
        # Use robust loader that ignores corrupted media parts
        with st.spinner("Loading document..."):
            doc = load_document_ignoring_bad_media(file_bytes)
        
        st.info("ℹ️ Note: Images are not analyzed. Only text formatting is checked.")
        
        all_issues = []
        for i, para in enumerate(doc.paragraphs):
            all_issues.extend(check_paragraph(i, para))

        if all_issues:
            st.success(f"Found {len(all_issues)} formatting issues.")
            st.dataframe(all_issues)

            # CSV download
            csv_output = io.StringIO()
            writer = csv.DictWriter(csv_output, fieldnames=["line", "type", "sentence", "suggestion"])
            writer.writeheader()
            writer.writerows(all_issues)
            csv_data = csv_output.getvalue()

            st.download_button(
                label="Download CSV Report",
                data=csv_data.encode('utf-8-sig'),
                file_name="formatting_issues.csv",
                mime="text/csv",
                key="csv_download"
            )

            # Text summary download
            text_summary = "\n".join([f"Line {issue['line']} | {issue['type']} | {issue['sentence']} | {issue['suggestion']}" for issue in all_issues])
            st.download_button(
                label="Download Text Summary",
                data=text_summary,
                file_name="formatting_issues.txt",
                mime="text/plain",
                key="txt_download"
            )
        else:
            st.info("No formatting issues found in the document!")
    except zipfile.BadZipFile as e:
        st.error(f"The uploaded file is not a valid DOCX (bad ZIP): {e}")
    except Exception as e:
        st.error(f"Error processing the file: {e}")
else:
    st.info("Please upload a .docx file to get started.")